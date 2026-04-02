import os
import json
import pandas as pd
from docx import Document
from datetime import datetime
import matplotlib.pyplot as plt
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# --- CONFIGURATION (Uses GitHub Secrets) ---
# Paste your Service Account JSON into a GitHub Secret named: GOOGLE_CREDENTIALS_JSON
SERVICE_ACCOUNT_INFO = os.getenv("GOOGLE_CREDENTIALS_JSON")
# The ID of the Google Drive folder for ROOT_OCEAN
FOLDER_ID = os.getenv("DRIVE_FOLDER_ID") 

def create_ocean_report(data, filename):
    """Generates a professional .docx with charts for ssgpt6.com submission."""
    doc = Document()
    doc.add_heading('GNAIAAAC LLC: ROOT OCEAN SUBMISSION', 0)
    
    # Process Data with Pandas
    df = pd.DataFrame(data)
    doc.add_heading('Data Summary', level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Create Visualization with Matplotlib
    plt.figure(figsize=(6, 4))
    plt.plot(df['timestamp'], df['metric'], marker='o', color='blue')
    plt.title('Ocean Metric Trends - Root Ocean')
    plt.grid(True)
    plot_path = 'ocean_trend.png'
    plt.savefig(plot_path)
    plt.close()
    
    doc.add_picture(plot_path)
    doc.add_paragraph("Analysis: The chart above indicates the projected ecosystem recovery metrics for the ROOT OCEAN project.")
    
    doc.save(filename)
    return filename

def upload_to_google_drive(file_path):
    """Uploads the file directly to the ROOT_OCEAN folder on Google Drive."""
    if not SERVICE_ACCOUNT_INFO:
        print("Error: GOOGLE_CREDENTIALS_JSON not found in Environment Variables.")
        return

    # Authenticate using the JSON secret
    info = json.loads(SERVICE_ACCOUNT_INFO)
    creds = service_account.Credentials.from_service_account_info(info)
    service = build('drive', 'v3', credentials=creds)

    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [FOLDER_ID] if FOLDER_ID else []
    }
    
    media = MediaFileUpload(file_path, resumable=True)
    uploaded_file = service.files().create(
        body=file_metadata, 
        media_body=media, 
        fields='id'
    ).execute()
    
    print(f"Success! File ID: {uploaded_file.get('id')} uploaded to Drive.")

if __name__ == "__main__":
    # Sample Data Payload
    ocean_data = {
        "timestamp": ["2026-01", "2026-02", "2026-03", "2026-04"],
        "metric": [45, 52, 61, 78]
    }
    
    report_name = "Root_Ocean_Claim_Analysis.docx"
    
    try:
        generated_file = create_ocean_report(ocean_data, report_name)
        upload_to_google_drive(generated_file)
    except Exception as e:
        print(f"Build Failed: {str(e)}")
        exit(1) # Ensures GitHub Actions registers the failure
